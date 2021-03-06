from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
import clipboard

#Open the document template
document =  Document("template.docx")

#Set and check filename
while True:
    filename = input("\nWhat should be the filename? Enter here:")
    check_filename = input("\nIs -- " + filename + " -- your filename? y/n: ")
    if filename != "template" and check_filename.lower() == "y"  or check_filename.lower() == "yes":
        break

#Set Question Style
styles = document.styles
set_question_style = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
question_style = document.styles['Question']
question_font = question_style.font
question_font.size = Pt(30)


#Set Answer Style 2
set_answer_style = styles.add_style("Answer 2", WD_STYLE_TYPE.PARAGRAPH)
answer_style = document.styles['Answer 2']
answer_font = answer_style.font
answer_font.size = Pt(42)

#Set Answer Style 3
set_answer_style = styles.add_style("Answer 3", WD_STYLE_TYPE.PARAGRAPH)
answer_style = document.styles['Answer 3']
answer_font = answer_style.font
answer_font.size = Pt(36)

#Set Answer Style 2
set_answer_style = styles.add_style("Answer 4", WD_STYLE_TYPE.PARAGRAPH)
answer_style = document.styles['Answer 4']
answer_font = answer_style.font
answer_font.size = Pt(24)



#Adds a main title to the first page of the doc
def add_title():
    while True:
        visual_title = input("\nEnter a title for today's visuals: ").capitalize()
        check_visual_title = input("\nIs  -- " + visual_title + " --  your title? y/n:")
        if check_visual_title.lower() == "y":
            break
    #Take the input and add it into the doc with styling
    visual_title_doc = document.add_paragraph(visual_title, style="Question")
    visual_title_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

#Checks clipboard and then adds the question to the doc
def add_question():
    print("\nAdd your question.")
    while True:
        question_title = input("\nEnter a Question: ").capitalize()
        check_question_title = input("\nIs  -- " + question_title + " --  your Question? y/n:")
        if check_question_title.lower() == "y":
            break
    question_title = document.add_paragraph(question_title, style="Question")
    question_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #Add spaces after question text
    for i in range(10):
        document.add_paragraph(" ")

#Add answer table
def add_answers():
    while True:
        try:
            num_questions = int(input("\nHow many answers are there for this question? Enter 2, 3, or 4: "))
            if num_questions == 2 or num_questions == 3 or num_questions == 4:
                break
        except ValueError:
            print("\nInput must be 2, 3, or 4")


    table = document.add_table(rows=1, cols=num_questions)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cells = table.rows[0].cells
    for i in range(len(cells)):
        print("\nAnswer " + str(i+1))
        while True:
            cell_title = input("\nEnter your Answer: ").capitalize()
            check_cell_title = input("\nIs  -- " + cell_title + " --  your title? y/n:")
            if check_cell_title.lower() == "y":
                break
        col = cells[i].add_paragraph(cell_title, style="Answer " + str(num_questions))
        col.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


#Main Loop
add_title()
while True:
    add_question()
    add_answers()
    finished_check = input("\n~~~~> Are you finished adding questions? <~~~~ y/n: ")
    if finished_check.lower() == "y" or finished_check.lower() == "yes":
        break

#Save the new document as filename and 
#Make sure that the filename is usable
while True:
    try:
        document.save(filename + ".docx")
        break
    except FileNotFoundError:
        filename = input("\nWhat should be the file name? Enter here: ")


test = input("\nworking?")